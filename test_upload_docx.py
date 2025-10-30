import requests
import os
from docx import Document

# Create a simple Word document for testing
def create_test_docx():
    doc = Document()
    doc.add_heading('Test Invoice', 0)
    doc.add_paragraph('Invoice Number: INV-001')
    doc.add_paragraph('Date: 2024-01-15')
    doc.add_paragraph('Customer: John Doe')
    doc.add_paragraph('Amount: $500.00')
    doc.add_paragraph('Description: Web Development Services')
    doc.save('test_invoice.docx')
    return 'test_invoice.docx'

# Test the analyze endpoint
url = "http://127.0.0.1:8000/analyze"

try:
    # Create test document
    filename = create_test_docx()
    
    # Test with the Word document
    with open(filename, 'rb') as f:
        files = {'file': (filename, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        response = requests.post(url, files=files)
    
    print(f"Status Code: {response.status_code}")
    print(f"Response: {response.text}")
    
    # Clean up
    os.remove(filename)
    
except Exception as e:
    print(f"Error: {e}")
    if hasattr(e, 'response') and e.response is not None:
        print(f"Response content: {e.response.text}")
